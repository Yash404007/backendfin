from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import logging
from datetime import datetime
from openai import AzureOpenAI
from PyPDF2 import PdfReader
import docx
from io import BytesIO
import re
from database import FinanceDocumentDB
import os
from dotenv import load_dotenv

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('finance_chatbot.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = FastAPI(title="Finance Document Q&A API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load environment variables from .env (if present)
load_dotenv()

# Configuration loaded from environment (use defaults if not provided)
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'database': os.getenv('DB_DATABASE', 'finance_docs'),
    'user': os.getenv('DB_USER', 'postgres'),
    'password': os.getenv('DB_PASSWORD', ''),
    'port': int(os.getenv('DB_PORT', 5432))
}

AZURE_CONFIG = {
    'endpoint': os.getenv('AZURE_ENDPOINT', ''),
    'api_key': os.getenv('AZURE_API_KEY', ''),
    'deployment': os.getenv('AZURE_DEPLOYMENT', ''),
    'api_version': os.getenv('AZURE_API_VERSION', '')
}

# Global instances
db = None
ai_client = None

# Pydantic models for request/response validation
class ChatRequest(BaseModel):
    document_id: Optional[int] = None
    document_ids: Optional[List[int]] = None
    message: str
    
class ChatResponse(BaseModel):
    response: str
    
class DocumentInfo(BaseModel):
    id: int
    document_name: str
    file_type: str
    word_count: int
    upload_date: datetime
    
class DocumentDetail(BaseModel):
    id: int
    document_name: str
    file_type: str
    content: str
    file_size: int
    word_count: int
    upload_date: datetime
    last_accessed: datetime
    
class ChatMessage(BaseModel):
    role: str
    content: str
    timestamp: datetime
    
class UploadResponse(BaseModel):
    document_id: int
    document_name: str
    word_count: int
    character_count: int
    
class DeleteResponse(BaseModel):
    message: str
    
class StatsResponse(BaseModel):
    total_documents: int
    total_messages: int
    total_content_size: int
    most_recent_document: Optional[str]
    most_recent_date: Optional[datetime]

# Helper functions for text extraction and cleaning
def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.replace('\x00', '')
    return text.strip()

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PdfReader(BytesIO(file_bytes))
        text_parts = []
        
        for i, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    logger.info(f"Extracted page {i+1}: {len(page_text)} characters")
            except Exception as e:
                logger.error(f"Error extracting page {i+1}: {str(e)}")
                continue
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"PDF extraction complete: {len(full_text)} characters")
        return full_text
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail=f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n".join(text_parts)
        logger.info(f"DOCX extraction complete: {len(full_text)} characters")
        return full_text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail=f"DOCX extraction failed: {str(e)}")

def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from TXT file with encoding detection"""
    try:
        try:
            text = file_bytes.decode('utf-8')
            logger.info("TXT decoded with UTF-8")
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
            logger.info("TXT decoded with latin-1")
        
        logger.info(f"TXT extraction complete: {len(text)} characters")
        return text
    except Exception as e:
        logger.error(f"TXT extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail=f"TXT extraction failed: {str(e)}")

# Application lifecycle events
@app.on_event("startup")
async def startup_event():
    """Initialize database and AI client on application startup"""
    global db, ai_client
    
    logger.info("Starting application initialization...")
    
    # Initialize database connection
    try:
        db = FinanceDocumentDB(**DB_CONFIG)
        if db.connect():
            db.create_tables()
            logger.info("Database connected and tables verified")
        else:
            logger.error("Failed to connect to database")
            raise Exception("Database connection failed")
    except Exception as e:
        logger.error(f"Database initialization failed: {str(e)}")
        raise
    
    # Initialize Azure OpenAI client
    try:
        ai_client = AzureOpenAI(
            azure_endpoint=AZURE_CONFIG['endpoint'],
            api_key=AZURE_CONFIG['api_key'],
            api_version=AZURE_CONFIG['api_version']
        )
        logger.info("Azure OpenAI client initialized successfully")
    except Exception as e:
        logger.error(f"Azure OpenAI initialization failed: {str(e)}")
        raise
    
    logger.info("Application initialization complete")

@app.on_event("shutdown")
async def shutdown_event():
    """Clean up resources on application shutdown"""
    if db:
        db.close()
        logger.info("Database connection closed")
    logger.info("Application shutdown complete")

# API Endpoints
@app.get("/")
async def root():
    """Root endpoint - API information"""
    return {
        "message": "Finance Document Q&A API",
        "version": "1.0.0",
        "status": "running",
        "endpoints": {
            "health": "/health",
            "docs": "/docs",
            "upload": "POST /documents/upload",
            "list_documents": "GET /documents",
            "chat": "POST /chat"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    db_status = db is not None and db.conn is not None
    ai_status = ai_client is not None
    
    overall_status = "healthy" if (db_status and ai_status) else "degraded"
    
    return {
        "status": overall_status,
        "database_connected": db_status,
        "ai_client_initialized": ai_status,
        "timestamp": datetime.now().isoformat()
    }

@app.post("/documents/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """
    Upload and process a finance document.
    Supports PDF, DOCX, and TXT file formats.
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    logger.info(f"Processing upload: {file.filename}")
    
    try:
        file_bytes = await file.read()
        
        # Determine file type and extract text
        content_type = file.content_type
        filename_lower = file.filename.lower()
        
        if content_type == "application/pdf" or filename_lower.endswith('.pdf'):
            text = extract_text_from_pdf(file_bytes)
            file_type = "pdf"
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename_lower.endswith('.docx'):
            text = extract_text_from_docx(file_bytes)
            file_type = "docx"
        elif content_type == "text/plain" or filename_lower.endswith('.txt'):
            text = extract_text_from_txt(file_bytes)
            file_type = "txt"
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file format: {content_type}. Supported formats: PDF, DOCX, TXT"
            )
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the file")
        
        # Clean the extracted text
        cleaned_text = clean_text(text)
        
        # Prepare metadata
        metadata = {
            "upload_timestamp": datetime.now().isoformat(),
            "file_size_bytes": len(file_bytes),
            "original_content_type": content_type
        }
        
        # Save to database
        document_id = db.save_document(
            document_name=file.filename,
            file_type=file_type,
            content=cleaned_text,
            file_size=len(file_bytes),
            metadata=metadata
        )
        
        word_count = len(cleaned_text.split())
        char_count = len(cleaned_text)
        
        logger.info(f"Document saved successfully: ID={document_id}, words={word_count}")
        
        return UploadResponse(
            document_id=document_id,
            document_name=file.filename,
            word_count=word_count,
            character_count=char_count
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload processing failed: {str(e)}")

@app.get("/documents", response_model=List[DocumentInfo])
async def get_documents():
    """
    Get list of all uploaded documents with metadata.
    Returns documents sorted by last accessed time (most recent first).
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    try:
        documents = db.get_all_documents()
        logger.info(f"Retrieved {len(documents)} documents")
        return documents
    except Exception as e:
        logger.error(f"Error retrieving documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve documents: {str(e)}")

@app.get("/documents/{document_id}")
async def get_document(document_id: int):
    """
    Get detailed information about a specific document including its content.
    This endpoint also updates the last_accessed timestamp.
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    try:
        document = db.get_document_by_id(document_id)
        if not document:
            raise HTTPException(status_code=404, detail=f"Document with ID {document_id} not found")
        
        logger.info(f"Retrieved document: ID={document_id}")
        return document
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error retrieving document {document_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve document: {str(e)}")

@app.delete("/documents/{document_id}", response_model=DeleteResponse)
async def delete_document(document_id: int):
    """
    Delete a document and all its associated chat history.
    This operation cannot be undone.
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    try:
        success = db.delete_document(document_id)
        if success:
            logger.info(f"Document deleted: ID={document_id}")
            return DeleteResponse(message=f"Document {document_id} deleted successfully")
        raise HTTPException(status_code=404, detail=f"Document with ID {document_id} not found")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """
    Chat with one or multiple documents using Azure OpenAI.
    
    For single document: provide document_id
    For cross-document analysis: provide document_ids (list of 2+ IDs)
    
    The AI will analyze document content and provide concise answers.
    Chat history is saved only for single document mode.
    """
    if not db or not ai_client:
        raise HTTPException(status_code=503, detail="Required services unavailable")
    
    logger.info(f"Chat request received: {request.message[:50]}...")
    
    try:
        # Validate request
        if not request.document_id and not request.document_ids:
            raise HTTPException(status_code=400, detail="Either document_id or document_ids must be provided")
        
        # Load document(s) and prepare system message
        if request.document_ids and len(request.document_ids) > 1:
            # Cross-document analysis mode
            logger.info(f"Cross-document mode: {len(request.document_ids)} documents")
            
            documents = db.get_multiple_documents(request.document_ids)
            if not documents:
                raise HTTPException(status_code=404, detail="One or more documents not found")
            
            # Combine document contents with clear separators
            combined_content = ""
            for doc in documents:
                combined_content += f"\n\n{'='*80}\n"
                combined_content += f"DOCUMENT: {doc['document_name']}\n"
                combined_content += f"{'='*80}\n\n"
                combined_content += doc['content']
            
            system_message = f"""You are a concise financial assistant analyzing MULTIPLE documents simultaneously. Answer questions by comparing, contrasting, and finding insights across all provided documents.

Document Contents:
{combined_content}

Critical Instructions:
1. CROSS-REFERENCE information across documents - identify similarities, differences, and patterns
2. CITE which document(s) you're referencing in your answer
3. REMEMBER the conversation context from previous questions
4. Be CONCISE and DIRECT - 3-5 sentences unless more detail is requested
5. When comparing, clearly state which document shows what
6. If information differs between documents, highlight the discrepancies
7. If the answer requires data from multiple documents, synthesize it clearly

Format your response with document citations and key insights."""
        
        elif request.document_id:
            # Single document mode
            logger.info(f"Single document mode: document_id={request.document_id}")
            
            document = db.get_document_by_id(request.document_id)
            if not document:
                raise HTTPException(status_code=404, detail=f"Document with ID {request.document_id} not found")
            
            system_message = f"""You are a concise financial assistant with conversation memory. Answer questions based ONLY on the following finance document content.

Document Content:
{document['content']}

Critical Instructions:
1. REMEMBER the conversation context - reference previous questions and answers when relevant
2. Be CONCISE and DIRECT - get straight to the point
3. Answer in 2-4 sentences maximum unless more detail is specifically requested
4. Use bullet points only when listing multiple items
5. Cite specific numbers, dates, or sections from the document
6. If the answer is not in the document, state: "This information is not available in the provided document"
7. After answering, suggest 2-3 intelligent follow-up questions based on the conversation flow"""
        
        else:
            raise HTTPException(status_code=400, detail="Invalid request: provide either document_id or multiple document_ids")
        
        # Get chat history (only for single document mode)
        conversation_history = []
        if request.document_id:
            chat_history = db.get_chat_history(request.document_id, limit=6)
            conversation_history = [
                {"role": msg['role'], "content": msg['content']} 
                for msg in chat_history
            ]
        
        # Prepare messages for Azure OpenAI API
        api_messages = [{"role": "system", "content": system_message}]
        api_messages.extend(conversation_history)
        api_messages.append({"role": "user", "content": request.message})
        
        logger.info(f"Sending request to Azure OpenAI (messages: {len(api_messages)})")
        
        # Call Azure OpenAI API
        response = ai_client.chat.completions.create(
            model=AZURE_CONFIG['deployment'],
            messages=api_messages,
            temperature=0.3,
            max_tokens=800
        )
        
        assistant_message = response.choices[0].message.content
        logger.info(f"Response received: {len(assistant_message)} characters")
        
        # Save chat history (only for single document mode)
        if request.document_id:
            db.save_chat_message(request.document_id, "user", request.message)
            db.save_chat_message(request.document_id, "assistant", assistant_message)
            logger.info(f"Chat history saved for document {request.document_id}")
        
        return ChatResponse(response=assistant_message)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Chat error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Chat processing failed: {str(e)}")

@app.get("/documents/{document_id}/history", response_model=List[ChatMessage])
async def get_chat_history(document_id: int, limit: int = 50):
    """
    Get chat history for a specific document.
    Returns up to 'limit' most recent messages in chronological order.
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    try:
        history = db.get_chat_history(document_id, limit)
        logger.info(f"Retrieved {len(history)} chat messages for document {document_id}")
        return history
    except Exception as e:
        logger.error(f"Error retrieving chat history: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve chat history: {str(e)}")

@app.delete("/documents/{document_id}/history", response_model=DeleteResponse)
async def clear_chat_history(document_id: int):
    """
    Clear all chat history for a specific document.
    The document itself is not deleted, only its conversation history.
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    try:
        # Verify document exists
        document = db.get_document_by_id(document_id)
        if not document:
            raise HTTPException(status_code=404, detail=f"Document with ID {document_id} not found")
        
        success = db.clear_chat_history(document_id)
        if success:
            logger.info(f"Chat history cleared for document {document_id}")
            return DeleteResponse(message=f"Chat history cleared for document {document_id}")
        raise HTTPException(status_code=500, detail="Failed to clear chat history")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error clearing chat history: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to clear chat history: {str(e)}")

@app.get("/stats", response_model=StatsResponse)
async def get_stats():
    """
    Get database statistics including total documents, messages, and storage usage.
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    try:
        stats = db.get_database_stats()
        logger.info("Database statistics retrieved")
        return stats
    except Exception as e:
        logger.error(f"Error getting stats: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve statistics: {str(e)}")

@app.get("/documents/search/{search_term}")
async def search_documents(search_term: str):
    """
    Search documents by name or content.
    Returns all documents that contain the search term.
    """
    if not db:
        raise HTTPException(status_code=503, detail="Database service unavailable")
    
    try:
        results = db.search_documents(search_term)
        logger.info(f"Search for '{search_term}' returned {len(results)} results")
        return {
            "search_term": search_term,
            "count": len(results),
            "results": results
        }
    except Exception as e:
        logger.error(f"Error searching documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

# Error handlers
@app.exception_handler(404)
async def not_found_handler(request, exc):
    return {"error": "Resource not found", "detail": str(exc)}

@app.exception_handler(500)
async def internal_error_handler(request, exc):
    logger.error(f"Internal server error: {str(exc)}")
    return {"error": "Internal server error", "detail": "An unexpected error occurred"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)